"""
=============================================================================
= addirap : générateur d'addition pour l'approche addirap
= auteur : Alix boc
= date : 2022
=============================================================================
"""
from pydoc import Doc
from random import randint
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.section import WD_SECTION,WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH,WD_BREAK
from docx.oxml.ns import qn
import sys
import datetime
import json

letter = {"width":21.51,"height":27.94}
half_letter = {"width":21.51,"height":13.97}
symboles = {"addition":"+","soustraction":"-"}

def getSymbole(operation):
    return symboles[operation]

def validateOperation(operation):
    if operation not in ['addition','soustraction']:
        print(operation + ": Paramètre invalide")
        sys.exit(1)
    return operation


def soustraction(minValue=1,maxValue=19):
    while True:
        a = randint(minValue,maxValue)
        b = randint(1,9)
        if (a-b) >= 0:
            return (a,b)
        elif (b-a) >= 0:
            return (b,a)

def addition(minValue=1,maxValue=9):
    a = randint(minValue,maxValue)
    b = randint(minValue,maxValue)
    return (a,b)

def generate(nbOperations=50,minValue=1,maxValue=9,operation="addition"):
    operations = []
    noOperation = 0
    maxLoop = 10000
    nbLoop = 0
    while noOperation < nbOperations:
        nbLoop += 1
        (a,b)=(-1,-1)
        if operation == "addition":
            (a,b) = addition(minValue,maxValue)
        elif operation == "soustraction":
            (a,b) = soustraction(minValue,maxValue)
        if nbLoop < maxLoop:
            if (a,b) not in operations:
                operations.append((a,b))
                noOperation += 1
        else:
            operations.append((a,b))
            noOperation += 1
    return operations

def formatDocxSection(doc):
    sections = doc.sections
    for section in sections:
        # section.orientation = WD_ORIENT.LANDSCAPE
        # section.page_width = Cm(letter["width"])
        # section.page_height = Cm(letter["height"])
        section.top_margin = Cm(0)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(1.5)
        section.right_margin = Cm(1.5)

def ajoutLigneDecoupe(doc):
    p1 = doc.add_paragraph("-"*122)
    p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

def saveDocx(mode,operations,operation,outputfile):
    doc = Document()
    for i in range(int(len(operations)/50)):
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        section.start_type

        sectPr = section._sectPr

        p1 = doc.add_heading("Niveau : " + mode + "\tOpération : " + operation + " \tdurée : ____ \tMon score : ____/" + str(len(operations[i*50:(i+1)*50])))
        p1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p1.add_run()
        run.add_break(WD_BREAK.LINE)
        section = doc.add_section(WD_SECTION.CONTINUOUS)
        section.start_type

        section = doc.sections[2]
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)

        # Set to 2 column layout
        sectPr2 = section._sectPr
        cols2 = sectPr.xpath('./w:cols')[0]
        cols2.set(qn('w:num'),'5')

        for key in operations[i*50:(i+1)*50]:
            p1 = doc.add_paragraph(str(key[0])+" "+symboles[operation]+" "+str(key[1])+" = __")

        section = doc.add_section(WD_SECTION.CONTINUOUS)
        section.start_type

        # Set to 2 column layout
        #sectPr = section._sectPr
        if i==0:
            cols = sectPr.xpath('./w:cols')[0]
            cols.set(qn('w:num'),'1')

    formatDocxSection(doc)
    doc.save(outputfile)

def saveJson(niveau,operations,operation,outputfile):
    result = {"operation":operation,"niveau":niveau,"operations":operations}
    json_object = json.dumps(result)
    out = open(outputfile,"w+")
    print(json_object,file=out)
    out.close()


def afficher(format,niveau,operations,operation="addition",outputfile="addirap.docx"):
    if format == "docx":
        saveDocx(niveau,operations,operation,outputfile)
    elif format == "json":
        saveJson(niveau,operations,operation,outputfile)

def validerNiveau(niveau,min,max):
    if niveau == "facile":
        minValue = min
        maxValue = int(max/2)+1
    elif niveau == "moyen":
        minValue = min
        maxValue = max
    elif niveau == "difficile":
        minValue = int(max/2)+1
        maxValue = max
    return (minValue,maxValue)

def main(parametres):
    nbOperations=50
    operation = "addition"
    minValue = -1
    maxValue = -1
    niveau = "moyen"
    format = "docx"
    if len(parametres) > 1:
        for param in parametres[1:]:
            tab = param.split("=")
            if tab[0] == "operation":
                operation = validateOperation(tab[1])
            elif tab[0] == "nbOperations":
                nbOperations = int(tab[1])
            elif tab[0] == "minValue":
                minValue = int(tab[1])
            elif tab[0] == "maxValue":
                maxValue = int(tab[1])
            elif tab[0] == "niveau":
                niveau = tab[1]
                (minValue,maxValue) = validerNiveau(tab[1],minValue,maxValue)
            elif tab[0] == "format":
                format = tab[1]
            else:
                print("parametre inconnu :",tab[0])
    
    #= valeur par défaut
    if operation == "soustraction" and maxValue==-1:
        maxValue = 19
    if minValue == -1:
        minValue=1
    if maxValue == -1:
        maxValue=9

    operations = generate(nbOperations=nbOperations,minValue=minValue,maxValue=maxValue,operation=operation)
    outputfile = "addirap-" + operation + "-" + niveau + "-" + datetime.datetime.now().strftime("%Y%m%d") + "." + format
    afficher(format,niveau,operations,operation=operation,outputfile=outputfile)

if __name__ == "__main__":
    main(sys.argv)
